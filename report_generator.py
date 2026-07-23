from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from pathlib import Path
from app.utils.logger import logger
from docx import Document

def generate_pdf_report(analysis):
    Path("reports").mkdir(exist_ok=True)
    pdf_path = f"reports/{analysis.contract_id}_report.pdf"
    
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=18)
    story.append(Paragraph(f"Contract Analysis: {analysis.filename}", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph(analysis.executive_summary or "", styles['Normal']))
    
    doc.build(story)
    logger.info(f"PDF generated: {pdf_path}")
    return pdf_path

def generate_word_report(analysis):
    Path("reports").mkdir(exist_ok=True)
    word_path = f"reports/{analysis.contract_id}_report.docx"
    doc = Document()
    doc.add_heading(f"Contract Analysis: {analysis.filename}", 0)
    doc.add_paragraph(analysis.executive_summary or "")
    doc.save(word_path)
    logger.info(f"Word report generated: {word_path}")
    return word_path